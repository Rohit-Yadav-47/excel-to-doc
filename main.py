import streamlit as st
import pandas as pd
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re  # Import regex module for robust parsing

# Define the Streamlit app
st.title("Excel to Word Converter")

st.write("Upload your Excel file to generate a Word document. Ensure the Excel file contains the following column names:")

# Display required column names
required_columns = [
    'Topic', 'Sub Topic', 'Difficulty Level', 'Question Text',
    'Choice 1', 'Choice 2', 'Choice 3', 'Choice 4', 'Correct choice',
    'Solution', 'Justification', 'Reference Link'
]
st.code("\n".join(required_columns))

# File uploader
uploaded_file = st.file_uploader("Upload Excel File", type=['xlsx'])

if uploaded_file:
    try:
        # Read the Excel file
        dataframe = pd.read_excel(uploaded_file)
        size = len(dataframe)  # Process all rows

        # Create Word document
        doc = docx.Document()

        def add_label_and_value(label, value):
            p = doc.add_paragraph()
            label_run = p.add_run(label)
            label_run.bold = True
            p.add_run(" " + value)

        for i in range(size):
            # Extract and clean data with fallback to empty string if NaN
            topic = str(dataframe.get('Topic', '').iloc[i]).strip() if 'Topic' in dataframe.columns else ""
            sub = str(dataframe.get('Sub Topic', '').iloc[i]).strip() if 'Sub Topic' in dataframe.columns else ""
            difficulty = str(dataframe.get('Difficulty Level', '').iloc[i]).strip() if 'Difficulty Level' in dataframe.columns else ""
            question = str(dataframe.get('Question Text', '').iloc[i]).strip() if 'Question Text' in dataframe.columns else ""
            choice1 = str(dataframe.get('Choice 1', '').iloc[i]).strip() if 'Choice 1' in dataframe.columns else ""
            choice2 = str(dataframe.get('Choice 2', '').iloc[i]).strip() if 'Choice 2' in dataframe.columns else ""
            choice3 = str(dataframe.get('Choice 3', '').iloc[i]).strip() if 'Choice 3' in dataframe.columns else ""
            choice4 = str(dataframe.get('Choice 4', '').iloc[i]).strip() if 'Choice 4' in dataframe.columns else ""
            correct_answer = str(dataframe.get('Correct choice', '').iloc[i]).strip() if 'Correct choice' in dataframe.columns else ""
            correct_sol = str(dataframe.get('Solution', '').iloc[i]).strip() if 'Solution' in dataframe.columns else ""
            justification = str(dataframe.get('Justification', '').iloc[i]).strip() if 'Justification' in dataframe.columns else ""
            reference_link = str(dataframe.get('Reference Link', '').iloc[i]).strip() if 'Reference Link' in dataframe.columns else ""

            # Add labels and values
            add_label_and_value("Topic -", topic)
            add_label_and_value("Sub Topic -", sub)
            add_label_and_value("Difficulty Level -", difficulty)

            # Add question
            q_para = doc.add_paragraph()
            q_para.add_run(f"Q{i + 1}. {question}")

            # Determine correct choice number using regex for robustness
            correct_num = None
            match = re.search(r'\d+', correct_answer.lower())
            if match:
                correct_num = match.group()

            def add_choice_line(num_str, choice_text, is_correct):
                p = doc.add_paragraph()
                # Add the choice number in bold
                num_run = p.add_run(f"({num_str}) ")
                num_run.bold = True

                # Add the choice text
                choice_run = p.add_run(choice_text)
                if is_correct:
                    choice_run.bold = True
                    choice_run.underline = True

            # Add choices
            if choice1:
                add_choice_line("1", choice1, is_correct=(correct_num == "1"))
            if choice2:
                add_choice_line("2", choice2, is_correct=(correct_num == "2"))
            if choice3:
                add_choice_line("3", choice3, is_correct=(correct_num == "3"))
            if choice4:
                add_choice_line("4", choice4, is_correct=(correct_num == "4"))

            # Add Correct Answer line fully in bold
            if correct_answer:
                ca_para = doc.add_paragraph()
                ca_run = ca_para.add_run("Correct Answer: " + correct_answer)
                ca_run.bold = True
            if correct_sol:
                sol_parat = doc.add_paragraph()
                sol_para = sol_parat.add_run("Solution:")
                sol_para.bold = True
                # Justification text in normal font
                for line in correct_sol.split("\n"):
                    doc.add_paragraph(line.strip())

            # Justification heading in bold
            if justification:
                just_para = doc.add_paragraph()
                just_heading = just_para.add_run("Justification:")
                just_heading.bold = True
                # Justification text in normal font
                for line in justification.split("\n"):
                    doc.add_paragraph(line.strip())

            # Reference Link in bold
            if reference_link:
                ref_para = doc.add_paragraph()
                ref_heading = ref_para.add_run("Reference Link:")
                ref_heading.bold = True
                # Reference link text on next line normal font
                doc.add_paragraph(reference_link)

            # Extra spacing before next entry
            doc.add_paragraph("")

        # Save the Word document
        output_file_name = uploaded_file.name.replace('.xlsx', '.docx')
        doc.save(output_file_name)

        # Provide the generated file for download
        with open(output_file_name, "rb") as f:
            st.download_button(
                label="Download Word Document",
                data=f,
                file_name=output_file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    except Exception as e:
        st.error(f"An error occurred: {e}")
